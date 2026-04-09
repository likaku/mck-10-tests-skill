#!/usr/bin/env python3
"""
McKinsey 10 Tests — 战略诊断报告生成器 v5
输出三种格式：长图 PNG · Word DOCX · PDF
白底黑字 · 楷体中文/Arial英文 · 雷达图 · 紧凑排版 · 大字号 · 自然语言
"""

import os
import math
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.font_manager as fm
import numpy as np
from PIL import Image, ImageDraw, ImageFont

# Word / PDF
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ── Fonts ─────────────────────────────────────────────────────
KAITI_PATH = '/System/Library/AssetsV2/com_apple_MobileAsset_Font8/88d6cc32a907955efa1d014207889413890573be.asset/AssetData/Kaiti.ttc'
ARIAL_PATH = '/Library/Fonts/Arial.ttf'
ARIAL_BOLD  = '/Library/Fonts/Arial Bold.ttf'

_font_cache = {}

def _get_font(path, size, index=0):
    key = (path, size, index)
    if key not in _font_cache:
        _font_cache[key] = ImageFont.truetype(path, size, index=index)
    return _font_cache[key]

def _kaiti(sz): return _get_font(KAITI_PATH, sz)
def _arial(sz): return _get_font(ARIAL_PATH, sz)
def _arial_b(sz):
    try: return _get_font(ARIAL_BOLD, sz)
    except: return _get_font(ARIAL_PATH, sz)

# ── Colors ────────────────────────────────────────────────────
BLK     = (15, 15, 15)        # 近黑
DK      = (30, 30, 30)        # 深灰（正文主色）
MD      = (70, 70, 70)        # 中灰
LT      = (160, 160, 160)
XLT     = (225, 225, 225)
BG      = (243, 243, 245)
WH      = (255, 255, 255)
NAVY    = (5, 28, 44)
RED     = (190, 42, 40)
ORANGE  = (195, 110, 25)
AMBER   = (155, 135, 30)
GREEN   = (32, 120, 65)

SCORE_CLR = {1: RED, 2: ORANGE, 3: AMBER, 4: GREEN}

# ── Layout ────────────────────────────────────────────────────
W       = 750
MG      = 40
CW      = W - 2 * MG

# ── Text engine ───────────────────────────────────────────────

def _wrap(text, font, max_w, draw):
    """Character-level wrap for CJK text. Returns list of lines."""
    if not text:
        return []
    lines = []
    for para in text.split('\n'):
        if not para.strip():
            lines.append('')
            continue
        cur = ''
        for ch in para:
            test = cur + ch
            bbox = draw.textbbox((0, 0), test, font=font)
            tw = bbox[2] - bbox[0]
            if tw > max_w and cur:
                lines.append(cur)
                cur = ch
            else:
                cur = test
        if cur:
            lines.append(cur)
    return lines


def _text(draw, x, y, text, font, fill=DK, max_w=None, leading=1.5):
    """Draw wrapped text, return new y."""
    if not text:
        return y
    if max_w is None:
        max_w = W - MG - x
    lines = _wrap(text, font, max_w, draw)
    lh = int(font.size * leading)
    for i, ln in enumerate(lines):
        draw.text((x, y + i * lh), ln, font=font, fill=fill)
    return y + len(lines) * lh


def _divider(draw, y, heavy=False):
    w = 2 if heavy else 1
    c = NAVY if heavy else XLT
    draw.line([(MG, y), (W - MG, y)], fill=c, width=w)
    return y + (10 if heavy else 6)


def _section(draw, y, title):
    """Big navy section header."""
    f = _kaiti(30)
    draw.text((MG, y), title, font=f, fill=NAVY)
    y += 42
    draw.line([(MG, y), (W - MG, y)], fill=NAVY, width=2)
    return y + 10


def _score_bar(draw, y, label, score, max_score=4):
    """Compact horizontal score bar."""
    bar_x = MG + 130
    bar_w = 220
    bar_h = 14
    clr = SCORE_CLR.get(score, MD)
    # label
    draw.text((MG, y), label, font=_kaiti(18), fill=DK)
    # bg bar
    draw.rounded_rectangle([(bar_x, y + 3), (bar_x + bar_w, y + 3 + bar_h)], radius=7, fill=XLT)
    # fill
    fw = max(4, int(bar_w * score / max_score))
    draw.rounded_rectangle([(bar_x, y + 3), (bar_x + fw, y + 3 + bar_h)], radius=7, fill=clr)
    # number
    draw.text((bar_x + bar_w + 10, y), f'{score}/{max_score}', font=_arial_b(17), fill=clr)
    return y + 30


# ── Radar chart ───────────────────────────────────────────────

def _radar(scores, labels, path):
    fm.fontManager.addfont(KAITI_PATH)
    plt.rcParams['font.family'] = ['Kaiti SC', 'Arial', 'sans-serif']
    N = len(labels)
    ang = np.linspace(0, 2 * np.pi, N, endpoint=False).tolist()
    vals = scores + [scores[0]]
    ang += [ang[0]]

    fig, ax = plt.subplots(figsize=(5.5, 5.5), subplot_kw=dict(polar=True))
    fig.patch.set_facecolor('white')
    ax.set_facecolor('white')
    ax.fill(ang, vals, color='#051C2C', alpha=0.07)
    ax.plot(ang, vals, color='#051C2C', linewidth=2.5)
    for a, s in zip(ang[:-1], scores):
        c = '#%02x%02x%02x' % SCORE_CLR[s]
        ax.scatter([a], [s], color=c, s=70, zorder=6, edgecolors='white', linewidths=1.5)
    ax.set_xticks(ang[:-1])
    ax.set_xticklabels(labels, fontsize=14, fontfamily='Kaiti SC', color='#222222')
    ax.set_ylim(0, 4.5)
    ax.set_yticks([1, 2, 3, 4])
    ax.set_yticklabels(['1', '2', '3', '4'], fontsize=10, color='#999999', fontfamily='Arial')
    ax.yaxis.grid(True, color='#E0E0E0', linewidth=0.6)
    ax.xaxis.grid(True, color='#E0E0E0', linewidth=0.6)
    ax.spines['polar'].set_visible(False)
    plt.tight_layout(pad=1)
    plt.savefig(path, dpi=160, bbox_inches='tight', facecolor='white', pad_inches=0.2)
    plt.close()


# ── Main ──────────────────────────────────────────────────────

def generate_report_card(data, output_path):
    radar_tmp = output_path.replace('.png', '_r.png')
    dims = data['dimensions']
    _radar([d['score'] for d in dims], [d['name'] for d in dims], radar_tmp)

    img = Image.new('RGB', (W, 9000), WH)
    draw = ImageDraw.Draw(img)
    y = 0

    # ━━ HEADER ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    y += 36
    draw.text((MG, y), data.get('title', 'McKinsey 10 Tests'), font=_kaiti(38), fill=NAVY)
    y += 52
    draw.text((MG, y), '战略诊断报告', font=_kaiti(22), fill=BLK)
    y += 34

    # Meta — 大字号、深色
    meta_font = _kaiti(16)
    meta_items = [
        f'诊断对象：{data["subject"]}',
        f'诊断顾问：{data["advisor"]}  ·  {data["date"]}',
    ]
    for line in meta_items:
        y = _text(draw, MG, y, line, meta_font, fill=MD)
        y += 2
    y += 8
    y = _divider(draw, y, heavy=True)

    # ━━ EXECUTIVE SUMMARY ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    y = _section(draw, y, '执行摘要')

    # Disclaimer
    draw.text((MG, y), '⚠ 本文档为初始想法的梳理，而非最终结论。', font=_kaiti(17), fill=ORANGE)
    y += 28

    # 综合评价 — 大字，黑色，密排
    y = _text(draw, MG, y, data['summary'], _kaiti(22), fill=BLK, leading=1.55)
    y += 12

    # Score pill
    score_text = f'  {data["total_score"]}/{data["max_score"]}  ·  {data["grade"]}  '
    badge_f = _arial_b(20)
    bbox = draw.textbbox((0, 0), score_text, font=badge_f)
    bw = bbox[2] - bbox[0] + 28
    bh = bbox[3] - bbox[1] + 16
    draw.rounded_rectangle([(MG, y), (MG + bw, y + bh)], radius=bh // 2, fill=NAVY)
    draw.text((MG + 14, y + 6), score_text, font=badge_f, fill=WH)
    y += bh + 18

    # ━━ KEY SHIFTS (自然语言叙述) ━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    y = _section(draw, y, '核心判断')

    shifts = data.get('shifts', [])
    for i, s in enumerate(shifts):
        # Number circle
        cx, cy_c = MG + 16, y + 16
        draw.ellipse([(cx - 15, cy_c - 15), (cx + 15, cy_c + 15)], fill=NAVY)
        draw.text((cx - 6, cy_c - 12), str(i + 1), font=_arial_b(20), fill=WH)

        # 标题
        y = _text(draw, MG + 44, y, s['heading'], _kaiti(22), fill=BLK, leading=1.5)
        y += 2
        # 正文
        y = _text(draw, MG + 44, y, s['body'], _kaiti(19), fill=DK, leading=1.55)
        y += 2
        # 理由（稍浅）
        if s.get('rationale'):
            y = _text(draw, MG + 44, y, s['rationale'], _kaiti(17), fill=MD, leading=1.5)
        y += 14

    # ━━ RADAR ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    y = _section(draw, y, '十维得分')

    ri = Image.open(radar_tmp)
    rw = CW
    rh = int(ri.height * rw / ri.width)
    ri = ri.resize((rw, rh), Image.LANCZOS)
    img.paste(ri, (MG, y))
    y += rh + 10

    # Score bars
    for d in dims:
        y = _score_bar(draw, y, d['name'], d['score'])
    y += 6

    # ━━ FINDINGS ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    y = _section(draw, y, '关键发现与建议')

    for d in dims:
        clr = SCORE_CLR.get(d['score'], MD)
        hdr = f'{d["name"]}  {d["score"]}/4'
        draw.text((MG, y), '●', font=_arial(16), fill=clr)
        draw.text((MG + 22, y), hdr, font=_kaiti(20), fill=BLK)
        y += 28
        # Finding — 大字深色
        y = _text(draw, MG + 10, y, d['finding'], _kaiti(18), fill=DK, max_w=CW - 10, leading=1.55)
        y += 3
        # Advice — 加粗 navy
        y = _text(draw, MG + 10, y, f'→ {d["advice"]}', _kaiti(18), fill=NAVY, max_w=CW - 10, leading=1.55)
        y += 14

    # ━━ STRENGTHS + IMPROVEMENTS ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    y = _section(draw, y, '亮点与改进')

    draw.text((MG, y), '最强 TOP 3', font=_kaiti(20), fill=GREEN)
    y += 28
    for i, s in enumerate(data['top3_strengths']):
        t = f'{i+1}. {s["name"]}（{s["score"]}/4）— {s["reason"]}'
        y = _text(draw, MG + 4, y, t, _kaiti(18), fill=DK, leading=1.55)
        y += 4
    y += 8

    draw.text((MG, y), '优先改进 TOP 3', font=_kaiti(20), fill=RED)
    y += 28
    for i, imp in enumerate(data['top3_improvements']):
        clr = RED if imp['score'] == 1 else ORANGE
        t = f'{i+1}. {imp["name"]}（{imp["score"]}/4）— {imp["reason"]}'
        y = _text(draw, MG + 4, y, t, _kaiti(18), fill=clr, leading=1.55)
        if imp.get('action'):
            y = _text(draw, MG + 18, y, f'→ {imp["action"]}', _kaiti(17), fill=NAVY, leading=1.5)
        y += 4
    y += 6

    # ━━ ACTIONS ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    y = _section(draw, y, '下一步行动')

    for a in data.get('priority_actions', []):
        draw.text((MG, y), a['icon'], font=_arial(20), fill=NAVY)
        draw.text((MG + 30, y), a['label'], font=_kaiti(20), fill=BLK)
        y += 28
        y = _text(draw, MG + 30, y, a['text'], _kaiti(18), fill=DK, leading=1.55)
        y += 10

    # ━━ FOOTER ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    y += 8
    y = _divider(draw, y)
    draw.text((MG, y), 'McKinsey Strategic 10 Tests · 初始想法梳理，非最终结论', font=_kaiti(14), fill=LT)
    y += 20
    draw.text((MG, y), f'Generated: {data["date"]}', font=_arial(12), fill=LT)
    y += 24

    # Crop & save
    img = img.crop((0, 0, W, y))
    os.makedirs(os.path.dirname(output_path) or '.', exist_ok=True)
    img.save(output_path, 'PNG', quality=95)
    if os.path.exists(radar_tmp):
        os.remove(radar_tmp)
    print(f'✅ {output_path} ({img.width}×{img.height}px)')
    return output_path


# ══════════════════════════════════════════════════════════════
# ██  WORD (.docx) GENERATOR  ██
# ══════════════════════════════════════════════════════════════

# Color helpers for docx
_NAVY_RGB   = RGBColor(5, 28, 44)
_BLK_RGB    = RGBColor(15, 15, 15)
_DK_RGB     = RGBColor(30, 30, 30)
_MD_RGB     = RGBColor(70, 70, 70)
_LT_RGB     = RGBColor(160, 160, 160)
_RED_RGB    = RGBColor(190, 42, 40)
_ORANGE_RGB = RGBColor(195, 110, 25)
_AMBER_RGB  = RGBColor(155, 135, 30)
_GREEN_RGB  = RGBColor(32, 120, 65)
_WHITE_RGB  = RGBColor(255, 255, 255)

_SCORE_RGB = {1: _RED_RGB, 2: _ORANGE_RGB, 3: _AMBER_RGB, 4: _GREEN_RGB}

KAITI_FONT_NAME = 'KaiTi'  # Word font name
ARIAL_FONT_NAME = 'Arial'


def _set_run(run, font_name=KAITI_FONT_NAME, size=11, color=_DK_RGB, bold=False):
    """Configure a run's font properties."""
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = font_name
    # Force East Asian font for CJK characters
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
        r.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}></w:rFonts>')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), KAITI_FONT_NAME)
    if font_name == ARIAL_FONT_NAME:
        rFonts.set(qn('w:ascii'), ARIAL_FONT_NAME)
        rFonts.set(qn('w:hAnsi'), ARIAL_FONT_NAME)


def _add_para(doc, text='', font_name=KAITI_FONT_NAME, size=11, color=_DK_RGB,
              bold=False, space_after=Pt(4), space_before=Pt(0), align=None):
    """Add a paragraph with configured styling."""
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = space_after
    pf.space_before = space_before
    pf.line_spacing = Pt(size * 1.5)
    if text:
        run = p.add_run(text)
        _set_run(run, font_name, size, color, bold)
    return p


def _add_navy_heading(doc, text, size=16):
    """Add a navy section heading with bottom border."""
    p = _add_para(doc, text, KAITI_FONT_NAME, size, _NAVY_RGB, bold=True,
                  space_before=Pt(14), space_after=Pt(2))
    # Bottom border
    pPr = p._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="12" w:space="3" w:color="051C2C"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    return p


def _add_divider(doc, heavy=False):
    """Add a horizontal line."""
    color = '051C2C' if heavy else 'E1E1E1'
    sz = '12' if heavy else '4'
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(2)
    pf.space_before = Pt(2)
    pPr = p._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="{sz}" w:space="1" w:color="{color}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    return p


def _shade_cell(cell, color_hex):
    """Apply background shading to a table cell."""
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>'
    )
    cell._element.get_or_add_tcPr().append(shading)


def generate_report_docx(data, output_path):
    """Generate a Word document matching the long-image report layout."""
    doc = Document()

    # ── Page setup: A4, narrow margins ──
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(1.5)

    # ── Default font ──
    style = doc.styles['Normal']
    font = style.font
    font.name = KAITI_FONT_NAME
    font.size = Pt(11)
    font.color.rgb = _DK_RGB
    rPr = style.element.find(qn('w:rPr'))
    if rPr is not None:
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}></w:rFonts>')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:eastAsia'), KAITI_FONT_NAME)

    dims = data['dimensions']

    # ━━ TITLE ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    _add_para(doc, data.get('title', 'McKinsey 10 Tests'),
              KAITI_FONT_NAME, 22, _NAVY_RGB, bold=True, space_after=Pt(2))
    _add_para(doc, '战略诊断报告', KAITI_FONT_NAME, 14, _BLK_RGB, bold=False, space_after=Pt(6))

    # Meta
    _add_para(doc, f'诊断对象：{data["subject"]}', KAITI_FONT_NAME, 10, _MD_RGB, space_after=Pt(1))
    _add_para(doc, f'诊断顾问：{data["advisor"]}  ·  {data["date"]}', KAITI_FONT_NAME, 10, _MD_RGB, space_after=Pt(4))
    _add_divider(doc, heavy=True)

    # ━━ EXECUTIVE SUMMARY ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    _add_navy_heading(doc, '执行摘要')

    # Disclaimer
    p = _add_para(doc, '', space_after=Pt(4))
    run = p.add_run('⚠ 本文档为初始想法的梳理，而非最终结论。')
    _set_run(run, KAITI_FONT_NAME, 10.5, _ORANGE_RGB, bold=False)

    # Summary body
    _add_para(doc, data['summary'], KAITI_FONT_NAME, 12, _BLK_RGB, space_after=Pt(6))

    # Score badge (table with navy bg)
    score_text = f'{data["total_score"]}/{data["max_score"]}  ·  {data["grade"]}'
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    _shade_cell(cell, '051C2C')
    p = cell.paragraphs[0]
    run = p.add_run(f'  {score_text}  ')
    _set_run(run, ARIAL_FONT_NAME, 13, _WHITE_RGB, bold=True)
    # Set cell width
    cell.width = Cm(6)
    _add_para(doc, '', space_after=Pt(4))

    # ━━ KEY SHIFTS ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    _add_navy_heading(doc, '核心判断')

    for i, s in enumerate(data.get('shifts', [])):
        # Numbered heading
        p = _add_para(doc, '', space_before=Pt(6), space_after=Pt(2))
        num_run = p.add_run(f'  {i+1}  ')
        _set_run(num_run, ARIAL_FONT_NAME, 11, _WHITE_RGB, bold=True)
        # Fake circle bg via highlight — use navy shading on the number
        from docx.oxml import OxmlElement
        rPr = num_run._element.get_or_add_rPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="051C2C"/>')
        rPr.append(shd)

        head_run = p.add_run(f'  {s["heading"]}')
        _set_run(head_run, KAITI_FONT_NAME, 13, _BLK_RGB, bold=True)

        # Body
        _add_para(doc, s['body'], KAITI_FONT_NAME, 11, _DK_RGB, space_after=Pt(2))

        # Rationale
        if s.get('rationale'):
            _add_para(doc, s['rationale'], KAITI_FONT_NAME, 10, _MD_RGB, space_after=Pt(6))

    # ━━ RADAR CHART ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    _add_navy_heading(doc, '十维得分')

    # Generate radar image
    radar_tmp = output_path.replace('.docx', '_radar.png')
    _radar([d['score'] for d in dims], [d['name'] for d in dims], radar_tmp)

    # Insert radar at document level (more reliable rendering)
    doc.add_picture(radar_tmp, width=Cm(14))
    # Center the last paragraph (which contains the picture)
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.paragraphs[-1].paragraph_format.space_after = Pt(8)

    # Score table with real progress bar images
    _bar_dir = os.path.join(os.path.dirname(output_path) or '.', '_bars')
    os.makedirs(_bar_dir, exist_ok=True)

    tbl = doc.add_table(rows=len(dims), cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Set table width to full page
    tbl_xml = tbl._tbl
    tblPr = tbl_xml.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}></w:tblPr>')
        tbl_xml.insert(0, tblPr)

    # Full-width table
    tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>')
    tblPr.append(tblW)

    for idx, d in enumerate(dims):
        sc = d['score']
        clr_rgb = SCORE_CLR.get(sc, MD)  # PIL color tuple
        clr_docx = _SCORE_RGB.get(sc, _MD_RGB)

        # Col 0: Dimension name
        cell0 = tbl.cell(idx, 0)
        cell0.width = Cm(3.5)
        p = cell0.paragraphs[0]
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.space_before = Pt(1)
        run = p.add_run(d['name'])
        _set_run(run, KAITI_FONT_NAME, 10, _DK_RGB)

        # Col 1: Progress bar as image
        cell1 = tbl.cell(idx, 1)
        cell1.width = Cm(9)
        # Generate a tiny progress bar image
        bar_w, bar_h = 400, 18
        bar_img = Image.new('RGB', (bar_w, bar_h), (240, 240, 240))
        bar_draw = ImageDraw.Draw(bar_img)
        # Background rounded rect
        bar_draw.rounded_rectangle([(0, 0), (bar_w - 1, bar_h - 1)], radius=bar_h // 2, fill=(225, 225, 225))
        # Filled portion
        fill_w = max(bar_h, int(bar_w * sc / 4))
        bar_draw.rounded_rectangle([(0, 0), (fill_w, bar_h - 1)], radius=bar_h // 2, fill=clr_rgb)
        bar_path = os.path.join(_bar_dir, f'bar_{idx}.png')
        bar_img.save(bar_path)

        p = cell1.paragraphs[0]
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.space_before = Pt(1)
        run = p.add_run()
        run.add_picture(bar_path, width=Cm(7), height=Cm(0.35))

        # Col 2: Score number
        cell2 = tbl.cell(idx, 2)
        cell2.width = Cm(2)
        p = cell2.paragraphs[0]
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.space_before = Pt(1)
        run = p.add_run(f'{sc}/4')
        _set_run(run, ARIAL_FONT_NAME, 10, clr_docx, bold=True)

    # Remove table borders for clean look
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

    _add_para(doc, '', space_after=Pt(4))

    # ━━ FINDINGS ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    _add_navy_heading(doc, '关键发现与建议')

    for d in dims:
        sc = d['score']
        clr = _SCORE_RGB.get(sc, _MD_RGB)

        # Dimension header
        p = _add_para(doc, '', space_before=Pt(4), space_after=Pt(1))
        dot = p.add_run('●  ')
        _set_run(dot, ARIAL_FONT_NAME, 10, clr)
        hdr = p.add_run(f'{d["name"]}  {sc}/4')
        _set_run(hdr, KAITI_FONT_NAME, 12, _BLK_RGB, bold=True)

        # Finding
        _add_para(doc, d['finding'], KAITI_FONT_NAME, 10.5, _DK_RGB, space_after=Pt(1))

        # Advice
        p = _add_para(doc, '', space_after=Pt(6))
        arrow = p.add_run('→ ')
        _set_run(arrow, ARIAL_FONT_NAME, 10.5, _NAVY_RGB, bold=True)
        adv = p.add_run(d['advice'])
        _set_run(adv, KAITI_FONT_NAME, 10.5, _NAVY_RGB)

    # ━━ STRENGTHS + IMPROVEMENTS ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    _add_navy_heading(doc, '亮点与改进')

    # Top 3
    p = _add_para(doc, '最强 TOP 3', KAITI_FONT_NAME, 12, _GREEN_RGB, bold=True, space_after=Pt(4))
    for i, s in enumerate(data['top3_strengths']):
        _add_para(doc, f'{i+1}. {s["name"]}（{s["score"]}/4）— {s["reason"]}',
                  KAITI_FONT_NAME, 10.5, _DK_RGB, space_after=Pt(2))

    _add_para(doc, '', space_after=Pt(4))

    # Weak 3
    p = _add_para(doc, '优先改进 TOP 3', KAITI_FONT_NAME, 12, _RED_RGB, bold=True, space_after=Pt(4))
    for i, imp in enumerate(data['top3_improvements']):
        clr = _RED_RGB if imp['score'] == 1 else _ORANGE_RGB
        _add_para(doc, f'{i+1}. {imp["name"]}（{imp["score"]}/4）— {imp["reason"]}',
                  KAITI_FONT_NAME, 10.5, clr, space_after=Pt(1))
        if imp.get('action'):
            p = _add_para(doc, '', space_after=Pt(4))
            arrow = p.add_run('    → ')
            _set_run(arrow, ARIAL_FONT_NAME, 10, _NAVY_RGB, bold=True)
            act = p.add_run(imp['action'])
            _set_run(act, KAITI_FONT_NAME, 10, _NAVY_RGB)

    # ━━ ACTIONS ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    _add_navy_heading(doc, '下一步行动')

    for a in data.get('priority_actions', []):
        p = _add_para(doc, '', space_before=Pt(4), space_after=Pt(1))
        icon_run = p.add_run(f'{a["icon"]}  ')
        _set_run(icon_run, ARIAL_FONT_NAME, 12, _NAVY_RGB)
        label_run = p.add_run(a['label'])
        _set_run(label_run, KAITI_FONT_NAME, 12, _BLK_RGB, bold=True)

        _add_para(doc, a['text'], KAITI_FONT_NAME, 10.5, _DK_RGB, space_after=Pt(6))

    # ━━ FOOTER ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    _add_divider(doc)
    _add_para(doc, 'McKinsey Strategic 10 Tests · 初始想法梳理，非最终结论',
              KAITI_FONT_NAME, 8, _LT_RGB, space_after=Pt(1))
    _add_para(doc, f'Generated: {data["date"]}  |  Apache 2.0 License  |  Kaku Li',
              ARIAL_FONT_NAME, 8, _LT_RGB, space_after=Pt(0))

    # Save
    os.makedirs(os.path.dirname(output_path) or '.', exist_ok=True)
    doc.save(output_path)

    # Cleanup temp files
    if os.path.exists(radar_tmp):
        os.remove(radar_tmp)
    # Cleanup bar images
    import shutil
    if os.path.exists(_bar_dir):
        shutil.rmtree(_bar_dir)

    print(f'✅ {output_path}')
    return output_path


def generate_report_pdf(data, docx_path, pdf_path=None):
    """Generate PDF from the Word document using docx2pdf (macOS: via Word/LibreOffice)."""
    if pdf_path is None:
        pdf_path = docx_path.replace('.docx', '.pdf')
    try:
        from docx2pdf import convert
        convert(docx_path, pdf_path)
        print(f'✅ {pdf_path}')
    except Exception as e:
        print(f'⚠️  PDF conversion failed: {e}')
        print('   Tip: docx2pdf requires Microsoft Word installed on macOS.')
        print(f'   The .docx file is ready at: {docx_path}')
    return pdf_path


def generate_all(data, base_path):
    """Generate all three formats: PNG long-image + DOCX + PDF."""
    png_path = base_path + '.png'
    docx_path = base_path + '.docx'
    pdf_path = base_path + '.pdf'

    generate_report_card(data, png_path)
    generate_report_docx(data, docx_path)
    generate_report_pdf(data, docx_path, pdf_path)

    return png_path, docx_path, pdf_path


# ── Demo ──────────────────────────────────────────────────────
if __name__ == '__main__':
    demo = {
        'title': 'McKinsey 10 Tests',
        'subject': '日世（中国）B端转型C端 / 出海东南亚',
        'scenario': '企业战略',
        'advisor': 'Prof. Sterling · 麦肯锡资深董事合伙人',
        'date': '2026-04-08',
        'total_score': 25, 'max_score': 40, 'grade': 'C+',
        'summary': (
            '日世拥有稀缺的品牌资产——日本软冰淇淋开创者、弟弟妹妹IP、设备+原料闭环——'
            '董事长的执行决心是最大亮点。但战略的核心问题在于聚焦不足和独到洞见缺失：'
            '同时打太多仗，又在最关键的战场（东南亚）缺乏信息差。'
            '这是一个"底子好但战略纪律不够"的典型案例。'
        ),
        'shifts': [
            {
                'heading': '从"求量"到"求质"',
                'body': '在中国B端，放弃挽回所有流失份额的幻想，聚焦3-5个高价值客户深度绑定，以利润率（>15%）而非销量作为核心KPI。',
                'rationale': '肯德基已在培养替代供应商——继续打价格战只会加速利润流失，不如主动收缩战线、保住利润池。',
            },
            {
                'heading': '从"五线并行"到"单点突破"',
                'body': '在东南亚第一年只做B端大客户供应这一件事。在中国，五条线砍掉至少两条。',
                'rationale': '资源有限时同时做五件事 = 每件事做到20%。先把一件事做到80%，再复制模式。',
            },
            {
                'heading': '从"同一起跑线"到"建立信息差"',
                'body': '董事长在投入第一分钱之前，花3个月深度浸泡东南亚目标市场——蹲便利店、跑经销商、理解冷链痛点，建立2-3个独家认知。',
                'rationale': '明治、格力高已在东南亚深耕多年。没有信息差就是在用自己的弱点打别人的强点。',
            },
        ],
        'dimensions': [
            {'name': '市场竞胜力', 'name_en': 'Beat the Market', 'score': 2,
             'finding': '中国B端正失去竞争力——大客户培养替代供应商，价格战压缩利润。C端盒马已失败，东南亚尚未验证。',
             'advice': '聚焦3-5个高价值客户，年利润率目标>15%，放弃"挽回所有份额"的思路'},
            {'name': '优势来源', 'name_en': 'Advantage Source', 'score': 3,
             'finding': '底层资产强——日本开创者品牌、弟弟妹妹IP、设备+原料闭环。但优势未被翻译成消费者可感知的差异化。',
             'advice': '在东南亚重新包装"日本品牌"故事，用创始人叙事+IP情感化做差异化核武器'},
            {'name': '精准聚焦', 'name_en': 'Granularity', 'score': 2,
             'finding': '同时在中国打五条线，东南亚又走三条路。资源严重分散，没有一条被充分验证。',
             'advice': '东南亚第一年只做B端大客户供应；中国砍掉或暂停至少两条线'},
            {'name': '趋势前瞻', 'name_en': 'Ahead of Trends', 'score': 3,
             'finding': '对大趋势判断正确——中国B端红海化、东南亚冷饮爆发。但看到机会≠做出取舍。',
             'advice': '每个新机会必须回答"为什么只有我能抓住"——答不上来的，不做'},
            {'name': '独到洞见', 'name_en': 'Privileged Insights', 'score': 1,
             'finding': '最大红色预警：在最关键的东南亚战场，与竞争对手站在同一起跑线——不了解当地口味、供应链、客户决策逻辑。',
             'advice': '董事长3个月深度浸泡东南亚（蹲便利店+跑经销商+理解冷链），建2-3个独家认知'},
            {'name': '不确定性管理', 'name_en': 'Uncertainty', 'score': 3,
             'finding': '有明确止损线（2027年6月/5000万），但缺少中间里程碑检查点。',
             'advice': '设3个季度检查点：Q4签首客→Q1首批交付→Q2月销目标，miss则立即复盘'},
            {'name': '承诺-灵活', 'name_en': 'Commit vs Flex', 'score': 3,
             'finding': '1000万+董事长亲自下场，承诺度够高。先B后C的节奏合理。',
             'advice': 'B端是"无悔之举"，C端/加盟是"期权"——等B端站稳后再激活'},
            {'name': '去偏见', 'name_en': 'Free from Bias', 'score': 2,
             'finding': '存在"选择逃避偏见"——每次面对取舍时"都要"而非"只选一个"。',
             'advice': '董事会引入红队机制：指定一人专门质疑每个新方向，强制"只能做一件事"的思考实验'},
            {'name': '执行决心', 'name_en': 'Conviction', 'score': 4,
             'finding': '最强维度——董事长亲自带队出海，愿意亏损1000万，有明确止损线。真正的战略承诺。',
             'advice': '保持决心，但警惕沉没成本：季度检查点miss时要有壮士断腕的勇气'},
            {'name': '行动计划', 'name_en': 'Action Plan', 'score': 3,
             'finding': '有计划和时间表，但尚在早期——还没有签下第一个东南亚客户来验证模式。',
             'advice': '第一个90天唯一目标：签下东南亚第一个标杆客户，用这一个case跑通全流程'},
        ],
        'top3_strengths': [
            {'name': '执行决心', 'score': 4, 'reason': '董事长亲自下场+1000万承诺，最稀缺的战略资源'},
            {'name': '优势来源', 'score': 3, 'reason': '日本开创者品牌+IP+技术闭环，底牌够硬'},
            {'name': '趋势前瞻', 'score': 3, 'reason': '对大趋势判断准确，方向没选错'},
        ],
        'top3_improvements': [
            {'name': '独到洞见', 'score': 1, 'reason': '最关键战场无信息差',
             'action': '董事长3个月深度浸泡东南亚，建2-3个独家认知'},
            {'name': '精准聚焦', 'score': 2, 'reason': '战线太长资源分散',
             'action': '东南亚第一年只做B端；中国砍掉≥2条线'},
            {'name': '去偏见', 'score': 2, 'reason': '"都要做"的逃避取舍偏见',
             'action': '引入红队机制，强制"只做一件事"实验'},
        ],
        'priority_actions': [
            {'icon': '🔥', 'label': '最紧急',
             'text': '董事长3个月内完成东南亚深度浸泡（至少2周驻扎），建立2-3个独家认知——所有后续行动的前提。'},
            {'icon': '💡', 'label': '最重要',
             'text': '在东南亚第一年只做B端大客户供应一件事；中国五条线砍掉至少两条——没有聚焦，再好的资源也是浪费。'},
            {'icon': '⚡', 'label': '48小时可做',
             'text': '列出中国五条线资源投入占比，标注哪两条可立即暂停；开始收集东南亚目标市场基础信息。'},
        ],
    }

    base = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'output', '日世战略诊断报告')
    generate_all(demo, base)
